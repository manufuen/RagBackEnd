import os
import re
from pathlib import Path

import docx
from pypdf import PdfReader
from sklearn.feature_extraction.text import CountVectorizer


SPANISH_STOPWORDS = [
    "de", "la", "que", "el", "en", "y", "a", "los", "del", "se",
    "las", "por", "un", "para", "con", "no", "una", "su", "al",
    "lo", "es", "como", "más", "pero", "sus", "le", "ya", "o",
    "este", "sí", "porque", "esta", "entre", "cuando", "muy",
    "sin", "sobre", "también", "me", "hasta", "hay", "donde",
    "quien", "desde", "todo", "nos", "durante", "todos", "uno",
    "les", "ni", "contra", "otros", "ese", "eso", "ante", "ellos",
    "e", "esto", "mí", "antes", "algunos", "qué", "unos", "yo",
    "otro", "otras", "otra", "él", "tanto", "esa", "estos", "mucho",
]


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    if ext == ".txt":
        return extract_text_from_txt(file_path)

    raise ValueError(f"Formato no soportado: {ext}")


def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)

    pages_text = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(page_text)

    return "\n".join(pages_text)


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    return "\n".join(paragraphs)


def extract_text_from_txt(file_path: str) -> str:
    return Path(file_path).read_text(
        encoding="utf-8",
        errors="ignore"
    )


def extract_author(file_path: str) -> str | None:
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            metadata = reader.metadata

            if metadata and metadata.author:
                return str(metadata.author)

        if ext == ".docx":
            document = docx.Document(file_path)
            author = document.core_properties.author

            if author:
                return str(author)

    except Exception:
        return None

    return None


def clean_text_for_keywords(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-záéíóúüñ0-9\s]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_keywords(text: str, max_keywords: int = 8) -> list[str]:
    cleaned_text = clean_text_for_keywords(text)

    if not cleaned_text:
        return []

    try:
        vectorizer = CountVectorizer(
            stop_words=SPANISH_STOPWORDS,
            max_features=max_keywords,
            ngram_range=(1, 2)
        )

        matrix = vectorizer.fit_transform([cleaned_text])
        keywords = vectorizer.get_feature_names_out().tolist()

        return keywords

    except Exception:
        return []